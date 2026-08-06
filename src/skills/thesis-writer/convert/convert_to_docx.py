"""
本科毕业论文 Markdown → Word 格式转换工具

用法:
    python convert_to_docx.py thesis.md output.docx

依赖:
    pip install -r requirements.txt
"""

import re
import os
import sys
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy


# ============================================================
# 全局配置
# ============================================================

FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
FONT_HEI = '黑体'
FONT_KAI = '楷体'

PAGE_CONFIG = {
    'top': Cm(3.0),
    'bottom': Cm(2.5),
    'left': Cm(2.5),
    'right': Cm(2.5),
    'gutter': Cm(0.5),
    'header_distance': Cm(1.5),
    'footer_distance': Cm(1.75),
}

HEADER_TEXT = ''  # 页眉文字，为空则不添加页眉


# ============================================================
# 工具函数
# ============================================================

def set_run_font(run, font_cn=FONT_CN, font_en=FONT_EN, size=Pt(12),
                 bold=False, italic=False):
    """设置 run 的中英文字体"""
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run.font.name = font_en
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


def set_paragraph_spacing(para, line_spacing=1.5, space_before=0, space_after=0):
    """设置段落间距"""
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def add_page_number(section, start=1):
    """添加页码（右对齐）"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    set_run_font(run, size=Pt(9))
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)


def add_header(section, text='', first_page_different=False):
    """添加页眉"""
    if not text:
        return
    section.different_first_page_header_footer = first_page_different
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, font_cn=FONT_CN, font_en=FONT_EN, size=Pt(9))


def create_paragraph(doc, text='', font_cn=FONT_CN, font_en=FONT_EN,
                     size=Pt(12), bold=False, alignment=None,
                     line_spacing=1.5, first_line_indent=None,
                     space_before=0, space_after=0):
    """创建格式化段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    set_paragraph_spacing(p, line_spacing, space_before, space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if text:
        run = p.add_run(text)
        set_run_font(run, font_cn, font_en, size, bold)
    return p


def add_chapter_title(doc, text, level=1):
    """添加章节标题"""
    if level == 1:
        p = create_paragraph(doc, text, FONT_HEI, FONT_EN, Pt(15),
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             line_spacing=1.5, space_before=12, space_after=6)
    elif level == 2:
        p = create_paragraph(doc, text, FONT_HEI, FONT_EN, Pt(14),
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             line_spacing=1.5, first_line_indent=Cm(0.74),
                             space_before=6, space_after=3)
    elif level == 3:
        p = create_paragraph(doc, text, FONT_CN, FONT_EN, Pt(12),
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             line_spacing=1.5, space_before=3, space_after=3)
    else:
        p = create_paragraph(doc, text, FONT_CN, FONT_EN, Pt(12),
                             bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             line_spacing=1.5, space_before=0, space_after=0)
    return p


def add_body_text(doc, text):
    """添加正文段落"""
    return create_paragraph(doc, text, FONT_CN, FONT_EN, Pt(12),
                            bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.5, first_line_indent=Cm(0.74))


def add_table_caption(doc, text):
    """添加表标题（表上方，居中，宋体五号加粗）"""
    p = create_paragraph(doc, text, FONT_CN, FONT_EN, Pt(10.5),
                         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.5, space_before=6, space_after=3)
    return p


def add_figure_caption(doc, text):
    """添加图标题（图下方，居中，宋体五号加粗）"""
    p = create_paragraph(doc, text, FONT_CN, FONT_EN, Pt(10.5),
                         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.5, space_before=3, space_after=12)
    return p


def add_empty_line(doc, count=1):
    """添加空行"""
    for _ in range(count):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 1.5, 0, 0)


def set_table_style(table):
    """设置三线表样式（顶线双线、表头下线单线、底线双线，无竖线，数据行之间无横线）"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="double" w:sz="6" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="double" w:sz="6" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    # 表头行下方加单线
    if table.rows:
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                '</w:tcBorders>'
            )
            tcPr.append(tcBorders)


def format_table_cell(cell, text='', font_cn=FONT_CN, font_en=FONT_EN,
                      size=Pt(10.5), bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """格式化表格单元格"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    set_paragraph_spacing(p, 1.5, 0, 0)
    run = p.add_run(text)
    set_run_font(run, font_cn, font_en, size, bold)


def add_three_line_table(doc, headers, rows, caption=''):
    """添加三线表"""
    if caption:
        add_table_caption(doc, caption)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True
    set_table_style(table)
    # 表头
    for i, h in enumerate(headers):
        format_table_cell(table.rows[0].cells[i], h, bold=True)
    # 数据行
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            format_table_cell(table.rows[r + 1].cells[c], str(cell_text), bold=False)
    add_empty_line(doc)
    return table


# ============================================================
# Markdown 解析
# ============================================================

def parse_markdown(md_text):
    """将 markdown 解析为结构化数据"""
    html = markdown.markdown(md_text, extensions=['tables', 'fenced_code', 'codehilite'])
    return BeautifulSoup(html, 'html.parser')


# ============================================================
# 主转换函数
# ============================================================

class ThesisConverter:
    def __init__(self, md_path):
        self.doc = Document()
        self.md_path = md_path
        self.md_dir = os.path.dirname(os.path.abspath(md_path))
        self.chapter_counters = {}
        self.current_chapter = 0
        self.figure_counter = 0
        self.table_counter = 0
        self.formula_counter = 0
        self.page_number_started = False

        with open(md_path, 'r', encoding='utf-8') as f:
            self.md_text = f.read()

        # 按 --- 分割章节
        self.sections = self._split_sections()

    def _split_sections(self):
        """按 --- 分割内容为章节块"""
        parts = re.split(r'\n---+\n', self.md_text)
        sections = []
        for part in parts:
            part = part.strip()
            if not part:
                continue
            # 解析 frontmatter
            fm = {}
            if part.startswith('---'):
                end = part.find('---', 3)
                if end > 0:
                    fm_text = part[3:end].strip()
                    for line in fm_text.split('\n'):
                        if ':' in line:
                            k, v = line.split(':', 1)
                            fm[k.strip()] = v.strip()
                    part = part[end + 3:].strip()
            sections.append({'frontmatter': fm, 'content': part})
        return sections

    def setup_page(self):
        """设置页面"""
        for section in self.doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = PAGE_CONFIG['top']
            section.bottom_margin = PAGE_CONFIG['bottom']
            section.left_margin = PAGE_CONFIG['left']
            section.right_margin = PAGE_CONFIG['right']
            section.gutter = PAGE_CONFIG['gutter']
            section.header_distance = PAGE_CONFIG['header_distance']
            section.footer_distance = PAGE_CONFIG['footer_distance']

    def add_abstract_cn(self, text, keywords):
        """添加中文摘要"""

        for para in text.split('\n'):
            para = para.strip()
            if para:
                add_body_text(self.doc, para)

        add_empty_line(self.doc)

        # 关键词
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, 1.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_label = p.add_run('关键词：')
        set_run_font(run_label, FONT_HEI, FONT_EN, Pt(12), bold=True)
        run_kw = p.add_run('；'.join(keywords))
        set_run_font(run_kw, FONT_CN, FONT_EN, Pt(12), bold=False)

        self.doc.add_page_break()

    def add_abstract_en(self, text, keywords):
        """添加英文摘要"""
        create_paragraph(self.doc, 'ABSTRACT', FONT_EN, FONT_EN, Pt(15),
                         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
        add_empty_line(self.doc)

        for para in text.split('\n'):
            para = para.strip()
            if para:
                add_body_text(self.doc, para)

        add_empty_line(self.doc)

        # Keywords
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, 1.5)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_label = p.add_run('Keywords：')
        set_run_font(run_label, FONT_EN, FONT_EN, Pt(12), bold=True)
        run_kw = p.add_run('; '.join(keywords))
        set_run_font(run_kw, FONT_EN, FONT_EN, Pt(12), bold=False)

        self.doc.add_page_break()

    def add_toc(self, toc_items):
        """添加目录（简化版）"""
        create_paragraph(self.doc, '目  录', FONT_HEI, FONT_EN, Pt(15),
                         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
        add_empty_line(self.doc)

        for item in toc_items:
            level = item.get('level', 1)
            title = item.get('title', '')
            page = item.get('page', '')

            p = self.doc.add_paragraph()
            set_paragraph_spacing(p, 1.5)

            if level == 1:
                # 一级目录：黑体小四加粗
                run = p.add_run(title)
                set_run_font(run, FONT_HEI, FONT_EN, Pt(12), bold=True)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                # 二级目录：黑体小四不加粗，左缩进
                run = p.add_run(title)
                set_run_font(run, FONT_HEI, FONT_EN, Pt(12), bold=False)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.left_indent = Cm(0.74)

            # 页码右对齐
            if page:
                tab_run = p.add_run(f'\t{page}')
                set_run_font(tab_run, FONT_EN, FONT_EN, Pt(12))

        self.doc.add_page_break()

    def convert_section(self, section_data):
        """转换一个章节"""
        content = section_data['content']
        soup = parse_markdown(content)
        self._process_elements(soup)

    def _handle_image(self, img_elem):
        alt = img_elem.get('alt', '')
        if alt:
            add_figure_caption(self.doc, alt)
        else:
            p = create_paragraph(self.doc, '[图片]', alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_paragraph_spacing(p, 1.5)

    def _process_elements(self, soup):
        """处理 HTML 元素"""
        for element in soup.children:
            if element.name is None:
                # 纯文本
                text = str(element).strip()
                if text:
                    add_body_text(self.doc, text)
                continue

            tag = element.name

            if tag in ['h1', 'h2', 'h3', 'h4', 'h5']:
                level = int(tag[1])
                text = element.get_text().strip()
                self.current_chapter = level if level <= 2 else self.current_chapter
                add_chapter_title(self.doc, text, level)

            elif tag == 'p':
                text = element.get_text().strip()
                # 检测是否为图表标题
                if re.match(r'(图|表)\d+[-.]?\d*', text):
                    if text.startswith('图'):
                        self.figure_counter += 1
                        add_figure_caption(self.doc, text)
                    elif text.startswith('表'):
                        self.table_counter += 1
                        add_table_caption(self.doc, text)
                # 检测段落中是否包含图片
                elif element.find('img'):
                    for img in element.find_all('img'):
                        self._handle_image(img)
                    if text:
                        add_body_text(self.doc, text)
                else:
                    for para in text.split('\n'):
                        para = para.strip()
                        if para:
                            add_body_text(self.doc, para)

            elif tag == 'table':
                self._convert_table(element)

            elif tag == 'ul' or tag == 'ol':
                for li in element.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    p = add_body_text(self.doc, text)

            elif tag == 'blockquote':
                text = element.get_text().strip()
                p = add_body_text(self.doc, text)
                p.paragraph_format.left_indent = Cm(1.5)

            elif tag == 'pre':
                code = element.get_text()
                p = add_body_text(self.doc, code)
                set_run_font(p.runs[0], FONT_EN, FONT_EN, Pt(10.5))

            elif tag == 'img':
                self._handle_image(element)

    def _convert_table(self, table_elem):
        """转换 HTML 表格为三线表"""
        headers = []
        rows_data = []
        caption = ''

        # 检查是否有 caption
        caption_elem = table_elem.find('caption')
        if caption_elem:
            caption = caption_elem.get_text().strip()

        # 提取表头
        thead = table_elem.find('thead')
        if thead:
            for th in thead.find_all('th'):
                headers.append(th.get_text().strip())
        else:
            first_row = table_elem.find('tr')
            if first_row:
                for th in first_row.find_all(['th', 'td']):
                    headers.append(th.get_text().strip())

        # 提取数据行
        tbody = table_elem.find('tbody') or table_elem
        start_row = 0 if thead else 1  # 如果没有 thead，跳过第一行
        for i, tr in enumerate(tbody.find_all('tr')):
            if i < start_row:
                continue
            row = []
            for td in tr.find_all('td'):
                row.append(td.get_text().strip())
            if row:
                rows_data.append(row)

        if headers:
            add_three_line_table(self.doc, headers, rows_data, caption)

    def convert(self, output_path):
        """执行完整转换"""
        self.setup_page()
        for section in self.doc.sections:
            add_header(section, HEADER_TEXT)
            add_page_number(section)

        for i, section in enumerate(self.sections):
            content = section['content'].strip()
            if not content:
                continue

            fm = section.get('frontmatter', {})

            # 处理特殊章节
            if fm.get('type') == 'abstract_cn':
                keywords = fm.get('keywords', '').replace('；', ';').split(';')
                self.add_abstract_cn(content, [k.strip() for k in keywords if k.strip()])
                continue
            elif fm.get('type') == 'abstract_en':
                keywords = fm.get('keywords', '').replace(';', ';').split(';')
                self.add_abstract_en(content, [k.strip() for k in keywords if k.strip()])
                continue
            elif fm.get('type') == 'toc':
                toc_items = []
                for line in content.split('\n'):
                    line = line.strip()
                    if line:
                        parts = line.split('|')
                        if len(parts) >= 2:
                            toc_items.append({
                                'level': int(parts[0].strip()),
                                'title': parts[1].strip(),
                                'page': parts[2].strip() if len(parts) > 2 else ''
                            })
                self.add_toc(toc_items)
                continue

            # 普通章节
            self.convert_section(section)

        # 保存
        self.doc.save(output_path)
        print(f'✅ 转换完成: {output_path}')


# ============================================================
# 命令行入口
# ============================================================

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('用法: python convert_to_docx.py thesis.md output.docx')
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    if not os.path.exists(input_path):
        print(f'错误: 找不到文件 {input_path}')
        sys.exit(1)

    converter = ThesisConverter(input_path)
    converter.convert(output_path)
